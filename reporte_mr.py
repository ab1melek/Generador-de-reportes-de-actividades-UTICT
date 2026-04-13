from __future__ import annotations

import calendar
import os
from collections import defaultdict
from datetime import datetime, date
from pathlib import Path

import requests
from dotenv import load_dotenv
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================================================
# CONFIGURACIÓN
# =========================================================
load_dotenv()


def get_env_str(name: str, default: str | None = None, required: bool = False) -> str:
    value = os.getenv(name, default)
    if required and not value:
        raise RuntimeError(f"Falta {name} en tu archivo .env")
    return value or ""


def get_env_int(name: str, default: int | None = None, required: bool = False) -> int | None:
    value = os.getenv(name)
    if value is None or value == "":
        return default
    try:
        return int(value)
    except ValueError as exc:
        raise RuntimeError(f"{name} debe ser un número entero") from exc

TOKEN = get_env_str("GITLAB_TOKEN", required=True)
BASE_URL = get_env_str("GITLAB_BASE_URL", "http://192.168.29.74:8091/api/v4")
AUTHOR_ID = get_env_int("GITLAB_AUTHOR_ID", required=True)
REPORT_YEAR = get_env_int("REPORT_YEAR", date.today().year)
REPORT_START_MONTH = get_env_int("REPORT_START_MONTH", 1, required=True)
REPORT_END_MONTH = get_env_int("REPORT_END_MONTH", REPORT_START_MONTH)

if REPORT_END_MONTH is None:
    REPORT_END_MONTH = REPORT_START_MONTH
if REPORT_END_MONTH < REPORT_START_MONTH:
    raise RuntimeError("REPORT_END_MONTH no puede ser menor que REPORT_START_MONTH")

START_DATE = date(REPORT_YEAR, REPORT_START_MONTH, 1)
END_DATE = date(
    REPORT_YEAR,
    REPORT_END_MONTH,
    calendar.monthrange(REPORT_YEAR, REPORT_END_MONTH)[1],
)

USER_NAME = get_env_str("USER_NAME", required=True)
ACTIVIDADES_CONTRATACION = get_env_str(
    "ACTIVIDADES_CONTRATACION",
    "DESARROLLAR APLICACIONES DE PRUEBAS END TO END Y UNITARIAS BASADAS EN "
    "VITEST SUPERTEST Y REACT TESTING LIBRARY PARA LOS SERVICIOS GENERADOS EN "
    "EL BACKEND Y FRONTEND ASÍ MISMO REALIZAR LA CREACIÓN DE DOCUMENTOS DE "
    "MEJORES PRÁCTICAS Y PROCESO APOYAR EN EL MATERIAL DIDÁCTICO Y WORKSHOP A "
    "EQUIPOS DE DESARROLLO DE LA UNIDAD DE TECNOLOGÍAS DE LA INFORMACIÓN Y "
    "COMUNICACIÓN DE LA TESORERÍA",
)

MONTH_NAMES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
}

HEADERS = {"PRIVATE-TOKEN": TOKEN}

SCRIPT_DIR = Path(__file__).resolve().parent
HEADER_PATH = SCRIPT_DIR / "header.png"
SIGNATURE_PATH = SCRIPT_DIR / "firma.png"
OUTPUT_DIR = Path.cwd()


# =========================================================
# UTILIDADES DE WORD
# =========================================================
def set_cell_text(cell, text: str, bold: bool = False, size: int = 11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def set_table_borders(table):
    """
    Aplica bordes simples a la tabla.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")

    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tblBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tblBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def hide_table_borders(table):
    """
    Oculta los bordes de una tabla para un aspecto más profesional.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")

    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tblBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tblBorders.append(element)
        element.set(qn("w:val"), "nil")


def add_header_image(doc: Document):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if HEADER_PATH.exists():
        run = p.add_run()
        run.add_picture(str(HEADER_PATH), width=Inches(6.7))
    else:
        # Si no existe el archivo, no truena el script.
        fallback = p.add_run("HEADER")
        fallback.bold = True


def setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    add_header_image(doc)


# =========================================================
# GITLAB
# =========================================================
def get_all_merge_requests() -> list[dict]:
    page = 1
    all_mr = []

    while True:
        url = (
            f"{BASE_URL}/merge_requests"
            f"?author_id={AUTHOR_ID}&scope=all&per_page=100&page={page}"
        )
        res = requests.get(url, headers=HEADERS, timeout=60)
        res.raise_for_status()
        data = res.json()

        if not data:
            break

        all_mr.extend(data)
        page += 1

    return all_mr


def parse_gitlab_datetime(value: str) -> datetime:
    # GitLab normalmente regresa formato ISO con Z
    return datetime.fromisoformat(value.replace("Z", "+00:00"))


def get_project_name(project_id: int, cache: dict[int, str]) -> str:
    if project_id in cache:
        return cache[project_id]

    url = f"{BASE_URL}/projects/{project_id}"
    res = requests.get(url, headers=HEADERS, timeout=60)

    if res.ok:
        data = res.json()
        # Preferimos path_with_namespace; si no, name
        name = data.get("path_with_namespace") or data.get("name") or str(project_id)
    else:
        name = str(project_id)

    cache[project_id] = name
    return name


def status_to_spanish(state: str) -> str:
    # Constraint: Estatus = En proceso o concluido
    return "Concluido" if state == "merged" else "En proceso"


def filter_by_period(mrs: list[dict]) -> list[dict]:
    filtered = []
    for mr in mrs:
        created = parse_gitlab_datetime(mr["created_at"]).date()
        if START_DATE <= created <= END_DATE:
            filtered.append(mr)
    return filtered


def group_by_month(mrs: list[dict]) -> dict[int, list[dict]]:
    grouped = defaultdict(list)
    for mr in mrs:
        created = parse_gitlab_datetime(mr["created_at"])
        grouped[created.month].append(mr)
    return grouped


# =========================================================
# GENERACIÓN DEL DOCX
# =========================================================
def create_monthly_report(month: int, mrs: list[dict], project_cache: dict[int, str]):
    month_name = MONTH_NAMES[month]
    last_day = calendar.monthrange(2026, month)[1]

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"reporte_{month_name}.docx"

    # Orden por fecha de creación
    mrs = sorted(mrs, key=lambda x: x["created_at"])

    doc = Document()
    setup_document(doc)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REPORTE DE ACTIVIDADES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Datos generales
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"Nombre: {USER_NAME}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"Periodo: 01/{month:02d}/2026 al {last_day:02d}/{month:02d}/2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"Actividades de Contratación: {ACTIVIDADES_CONTRATACION}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    doc.add_paragraph("")

    # Tabla

    doc.add_paragraph("")

    # Tabla
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)

    headers = ["No.", "PRODUCTO", "ACTIVIDADES PRINCIPALES", "ESTATUS"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=11)

    for idx, mr in enumerate(mrs, start=1):
        project_name = get_project_name(mr["project_id"], project_cache)
        activity_title = mr.get("title", "").strip()
        created_date = parse_gitlab_datetime(mr["created_at"]).date().isoformat()
        merged_date = mr.get("merged_at")
        merged_date_txt = ""
        if merged_date:
            merged_date_txt = parse_gitlab_datetime(merged_date).date().isoformat()

        # Aquí dejamos el contenido más cercano al formato original,
        # pero agregando fecha dentro de la actividad para que el reporte
        # quede más útil y trazable.
        activity_text = activity_title
        if created_date:
            activity_text = f"{activity_title}\nFecha: {created_date}"
        if merged_date_txt:
            activity_text += f"\nFusión: {merged_date_txt}"

        row = table.add_row().cells
        set_cell_text(row[0], str(idx))
        set_cell_text(row[1], project_name)
        set_cell_text(row[2], activity_text)
        set_cell_text(row[3], status_to_spanish(mr.get("state", "")))

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Firmas
    tabla_firmas = doc.add_table(rows=4, cols=2)
    tabla_firmas.style = "Table Grid"
    hide_table_borders(tabla_firmas)

    def set_firma(cell, texto="", bold=False):
        if not cell.paragraphs:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
            p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if texto:
            run = p.add_run(texto)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    def add_image_paragraph(cell, image_path, width=Inches(2.2)):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=width)

    top_cell = tabla_firmas.rows[0].cells[0].merge(tabla_firmas.rows[0].cells[1])
    top_cell.text = ""
    set_firma(top_cell, "_____________________________________")
    if SIGNATURE_PATH.exists():
        add_image_paragraph(top_cell, SIGNATURE_PATH, width=Inches(2.2))
    set_firma(top_cell, "Prestador de Servicios")
    set_firma(top_cell, f"Prestador de Servicios\n{USER_NAME}\nRealizó")

    set_firma(tabla_firmas.rows[1].cells[0], "")
    set_firma(tabla_firmas.rows[1].cells[1], "")

    set_firma(tabla_firmas.rows[2].cells[0], "_____________________________________")
    set_firma(tabla_firmas.rows[2].cells[1], "_________________________________________")

    set_firma(
        tabla_firmas.rows[3].cells[0],
        "Lic. Arturo Martínez Alvarado\nCoordinador de Desarrollo de Sistemas de la UTICT\nSupervisó",
    )
    set_firma(
        tabla_firmas.rows[3].cells[1],
        "Mtro. Luis Arturo López Caballero\nDirector de la UTICT\nAutorizó",
    )

    doc.save(output_path)
    print(f"Generado: {output_path}")


def main():
    print("Obteniendo merge requests...")
    all_mr = get_all_merge_requests()
    print(f"Total obtenidos: {len(all_mr)}")

    filtered = filter_by_period(all_mr)
    print(f"Filtrados {START_DATE.isoformat()} a {END_DATE.isoformat()}: {len(filtered)}")

    grouped = group_by_month(filtered)
    project_cache: dict[int, str] = {}

    for month in range(START_DATE.month, END_DATE.month + 1):
        month_mrs = grouped.get(month, [])
        if month_mrs:
            create_monthly_report(month, month_mrs, project_cache)
        else:
            print(f"Sin datos para {MONTH_NAMES.get(month, str(month))}")

if __name__ == "__main__":
    main()